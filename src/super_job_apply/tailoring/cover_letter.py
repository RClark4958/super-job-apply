"""AI-powered cover letter generation per job posting.

Runs the multi-agent writing pipeline (Writer → Editor → Mediator) to
produce a human-quality cover letter, logs all stages to audit, and
outputs a .docx file.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

from ..audit import AuditTrail
from ..models import CandidateProfile, JobPosting, MatchScore
from ..writers.pipeline import write_cover_letter

logger = logging.getLogger(__name__)


async def generate_cover_letter(
    job: JobPosting,
    candidate: CandidateProfile,
    score: MatchScore,
    output_dir: str,
    audit: AuditTrail,
    application_id: str | None = None,
    model: str = "claude-sonnet-4-6",
) -> tuple[str, str]:
    """Generate a tailored cover letter using the multi-agent writing pipeline.

    Args:
        job: The job posting.
        candidate: Candidate profile.
        score: Match score with keyword analysis.
        output_dir: Directory to save the cover letter.
        audit: Audit trail for logging.
        application_id: Optional application ID for audit linkage.
        model: Anthropic model for the writing agents.

    Returns:
        Tuple of (path to .docx file, cover letter body text).
    """
    # Run the Writer → Editor → Mediator pipeline
    letter_body = await write_cover_letter(
        job=job,
        candidate=candidate,
        score=score,
        audit=audit,
        application_id=application_id,
        model=model,
    )

    # Build full letter with header and sign-off
    full_letter = f"Dear Hiring Manager,\n\n{letter_body}\n\nSincerely,\n{candidate.name}"

    # Save as .docx
    output_path = _build_output_path(job, output_dir)
    _write_cover_letter_docx(full_letter, output_path)

    logger.info(f"Cover letter saved: {output_path}")
    return output_path, letter_body


def _build_output_path(job: JobPosting, output_dir: str) -> str:
    """Build a safe output file path for the cover letter."""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    safe_company = "".join(c if c.isalnum() or c in " -_" else "" for c in job.company_name)
    safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in job.job_title)
    filename = f"{safe_company}_{safe_title}_cover_letter.docx".replace(" ", "_")
    return str(Path(output_dir) / filename)


def _write_cover_letter_docx(text: str, output_path: str) -> None:
    """Write cover letter text into a clean .docx file."""
    doc = Document()

    for paragraph in text.split("\n"):
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)
        else:
            doc.add_paragraph("")

    doc.save(output_path)
