"""AI-powered resume tailoring per job posting.

Reads a base resume template (.docx), runs the multi-agent writing pipeline
(Writer → Editor → Mediator) to rewrite content, logs all stages to audit,
and outputs a tailored .docx file.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

from ..audit import AuditEventType, AuditTrail
from ..models import CandidateProfile, JobPosting, MatchScore
from ..writers.pipeline import write_resume

logger = logging.getLogger(__name__)


async def tailor_resume(
    job: JobPosting,
    candidate: CandidateProfile,
    score: MatchScore,
    template_path: str,
    output_dir: str,
    audit: AuditTrail,
    application_id: str | None = None,
    model: str = "claude-sonnet-4-6",
) -> str:
    """Generate a tailored resume using the multi-agent writing pipeline.

    Args:
        job: The job posting to tailor for.
        candidate: Candidate profile.
        score: Match score with keyword analysis.
        template_path: Path to base resume .docx template.
        output_dir: Directory to save tailored resume.
        audit: Audit trail for logging.
        application_id: Optional application ID for audit linkage.
        model: Anthropic model for the writing agents.

    Returns:
        Path to the generated tailored resume .docx file.
    """
    # Extract original content and log it
    original_content = _extract_docx_content(template_path)

    await audit.log_writing_stage(
        job_id=job.id,
        application_id=application_id,
        event_type=AuditEventType.RESUME_ORIGINAL,
        content=original_content,
    )

    # Run the Writer → Editor → Mediator pipeline
    final_content = await write_resume(
        job=job,
        candidate=candidate,
        score=score,
        original_content=original_content,
        audit=audit,
        application_id=application_id,
        model=model,
    )

    # Generate output .docx
    output_path = _build_output_path(job, output_dir, "resume")
    _write_tailored_docx(final_content, template_path, output_path)

    logger.info(f"Tailored resume saved: {output_path}")
    return output_path


def _extract_docx_content(path: str) -> str:
    """Extract text content from a .docx file."""
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style and para.style.name and "Heading" in para.style.name:
                lines.append(f"## {text}")
            elif text.startswith(("•", "-", "–", "◦")):
                lines.append(f"- {text.lstrip('•-–◦ ')}")
            else:
                lines.append(text)
    return "\n".join(lines)


def _build_output_path(job: JobPosting, output_dir: str, doc_type: str) -> str:
    """Build a safe output file path."""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    safe_company = "".join(c if c.isalnum() or c in " -_" else "" for c in job.company_name)
    safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in job.job_title)
    filename = f"{safe_company}_{safe_title}_{doc_type}.docx".replace(" ", "_")
    return str(Path(output_dir) / filename)


def _write_tailored_docx(content: str, template_path: str, output_path: str) -> None:
    """Write tailored content into a new .docx file based on the template's styling."""
    doc = Document(template_path)

    # Clear existing content
    for para in doc.paragraphs:
        para.clear()

    while len(doc.paragraphs) > 1:
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    style_names = [s.name for s in doc.styles]

    first = True
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue

        if first:
            para = doc.paragraphs[0]
            first = False
        else:
            para = doc.add_paragraph()

        if line.startswith("## "):
            para.text = line[3:]
            para.style = doc.styles["Heading 2"] if "Heading 2" in style_names else doc.styles["Heading 1"]
        elif line.startswith("# "):
            para.text = line[2:]
            para.style = doc.styles["Heading 1"]
        elif line.startswith("- "):
            para.text = line[2:]
            para.style = doc.styles["List Bullet"] if "List Bullet" in style_names else doc.styles["Normal"]
        else:
            para.text = line

    doc.save(output_path)
