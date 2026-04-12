"""CLI interface for super-job-apply."""

from __future__ import annotations

import asyncio
import logging
import sys

import click
from rich.console import Console
from rich.logging import RichHandler

from .config import load_config
from .db import Database
from .models import ApplicationStatus

console = Console()


def _setup_logging(verbose: bool = False) -> None:
    """Configure logging with rich handler."""
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(message)s",
        datefmt="[%X]",
        handlers=[RichHandler(console=console, rich_tracebacks=True)],
    )


@click.group()
@click.option("--verbose", "-v", is_flag=True, help="Enable debug logging")
def main(verbose: bool) -> None:
    """super-job-apply: AI-powered job application automation at scale."""
    _setup_logging(verbose)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
@click.option("--dry-run", is_flag=True, help="Score and tailor but don't submit applications")
def run(config_path: str, dry_run: bool) -> None:
    """Run the full job application pipeline."""
    try:
        config = load_config(config_path)
        if dry_run:
            config.application.dry_run = True

        from .pipeline import run_pipeline

        asyncio.run(run_pipeline(config))
    except FileNotFoundError as e:
        console.print(f"[red]{e}[/red]")
        sys.exit(1)
    except ValueError as e:
        console.print(f"[red]Configuration error: {e}[/red]")
        sys.exit(1)
    except KeyboardInterrupt:
        console.print("\n[yellow]Interrupted.[/yellow]")
        sys.exit(0)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def stats(config_path: str) -> None:
    """Show application statistics and recommendations."""
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)

        from .reporting.stats import show_stats

        asyncio.run(show_stats(db))
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
@click.option("--count", "-n", default=10, help="Number of recent applications to show")
def recent(config_path: str, count: int) -> None:
    """Show recent applications."""
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)

        from .reporting.stats import show_recent

        asyncio.run(show_recent(db, count))
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.argument("app_id")
@click.option(
    "--status",
    type=click.Choice(["interview", "rejected", "offer"], case_sensitive=False),
    required=True,
    help="New status for the application",
)
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def update(app_id: str, status: str, config_path: str) -> None:
    """Manually update an application status (e.g., mark as interview/rejected/offer)."""
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)

        status_enum = ApplicationStatus(status.lower())

        async def _update():
            await db.update_application(app_id, status=status_enum)

        asyncio.run(_update())
        console.print(f"[green]Application {app_id} updated to '{status}'[/green]")
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
@click.option(
    "--format", "fmt",
    type=click.Choice(["csv", "json"], case_sensitive=False),
    default="csv",
    help="Export format",
)
@click.option("--output", "-o", default=None, help="Output file path (default: stdout)")
def export(config_path: str, fmt: str, output: str | None) -> None:
    """Export application data as CSV or JSON."""
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)

        from .reporting.stats import export_data

        data = asyncio.run(export_data(db, fmt))

        if output:
            with open(output, "w") as f:
                f.write(data)
            console.print(f"[green]Exported to {output}[/green]")
        else:
            click.echo(data)
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
@click.option("--all-statuses", is_flag=True, help="Show all jobs, not just approved/pending/scored")
def submittable(config_path: str, all_statuses: bool) -> None:
    """List jobs that can be submitted (non-blocked ATS platforms).

    Shows jobs from the database that are NOT on blocked ATS platforms
    (Greenhouse, Lever) and have a submittable status.
    """
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)

        from .ats_detection import BLOCKED_ATS, detect_ats_platform, get_form_hints

        async def _list():
            await db.init()
            import aiosqlite

            async with aiosqlite.connect(config.application.db_path) as conn:
                conn.row_factory = aiosqlite.Row

                # Get all jobs joined with their latest application status
                query = """
                    SELECT j.id, j.company_name, j.job_title, j.careers_url,
                           j.location, j.work_type,
                           a.status, a.match_score, a.id as app_id
                    FROM jobs j
                    LEFT JOIN applications a ON a.job_id = j.id
                    ORDER BY a.match_score DESC NULLS LAST
                """
                cursor = await conn.execute(query)
                rows = await cursor.fetchall()

            submittable_jobs = []
            blocked_jobs = []
            platform_counts: dict[str, int] = {}

            for row in rows:
                url = row["careers_url"] or ""
                ats = detect_ats_platform(url)
                platform_counts[ats or "unknown"] = platform_counts.get(ats or "unknown", 0) + 1

                status = row["status"] or "no_application"
                if not all_statuses and status not in ("approved", "pending", "scored", "no_application"):
                    continue

                entry = {
                    "id": row["id"],
                    "company": row["company_name"],
                    "title": row["job_title"],
                    "url": url,
                    "ats": ats or "unknown",
                    "status": status,
                    "score": row["match_score"],
                    "app_id": row["app_id"],
                }

                if ats in BLOCKED_ATS:
                    blocked_jobs.append(entry)
                else:
                    submittable_jobs.append(entry)

            # Platform breakdown
            console.print("\n[bold]ATS Platform Breakdown (all jobs in DB):[/bold]")
            for platform, count in sorted(platform_counts.items(), key=lambda x: -x[1]):
                blocked_marker = " [red](BLOCKED)[/red]" if platform in BLOCKED_ATS else ""
                console.print(f"  {platform}: {count}{blocked_marker}")

            # Submittable jobs
            console.print(f"\n[bold green]Submittable Jobs ({len(submittable_jobs)}):[/bold green]")
            if submittable_jobs:
                for i, job in enumerate(submittable_jobs, 1):
                    score_str = f"{job['score']:.2f}" if job["score"] else "N/A"
                    console.print(
                        f"  {i:3}. [{job['ats']:16}] {job['company']:25} — {job['title'][:45]}"
                    )
                    console.print(
                        f"       Score: {score_str} | Status: {job['status']} | {job['url'][:70]}"
                    )
            else:
                console.print("  [yellow]No submittable jobs found.[/yellow]")

            # Blocked summary
            if blocked_jobs:
                console.print(
                    f"\n[dim]Blocked jobs ({len(blocked_jobs)}): "
                    f"on {', '.join(sorted(BLOCKED_ATS))} — skipped automatically[/dim]"
                )

            # Actionable summary
            ready = [j for j in submittable_jobs if j["status"] in ("approved",)]
            pending_review = [j for j in submittable_jobs if j["status"] in ("pending",)]
            no_app = [j for j in submittable_jobs if j["status"] == "no_application"]

            console.print(f"\n[bold]Summary:[/bold]")
            console.print(f"  Ready to submit (approved): {len(ready)}")
            console.print(f"  Needs review (pending):     {len(pending_review)}")
            console.print(f"  No application yet:         {len(no_app)}")
            console.print(f"  Blocked (captcha):          {len(blocked_jobs)}")

            if ready:
                console.print("\nRun [bold]super-job-apply submit[/bold] to submit approved jobs.")
            elif pending_review:
                console.print("\nRun [bold]super-job-apply review[/bold] to approve pending jobs.")

        asyncio.run(_list())
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.argument("job_id")
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def audit(job_id: str, config_path: str) -> None:
    """Show the full audit trail for a job (all writing stages, drafts, edits)."""
    try:
        config = load_config(config_path)

        from .audit import AuditTrail

        trail = AuditTrail(
            db_path=config.application.db_path,
            output_dir=config.application.output_dir,
        )

        async def _show():
            events = await trail.get_events_for_job(job_id)
            if not events:
                console.print(f"[yellow]No audit events found for job {job_id}[/yellow]")
                return

            console.print(f"\n[bold]Audit Trail for job {job_id}[/bold]")
            console.print("=" * 60)

            for event in events:
                event_type = event["event_type"]
                created = event["created_at"]
                if isinstance(created, str) and "T" in created:
                    created = created.split("T")[0] + " " + created.split("T")[1][:8]

                # Color by event type
                type_colors = {
                    "job_discovered": "blue",
                    "job_scored": "cyan",
                    "resume_original": "dim",
                    "resume_writer_draft": "yellow",
                    "resume_editor_review": "magenta",
                    "resume_mediator_final": "green",
                    "cover_letter_writer_draft": "yellow",
                    "cover_letter_editor_review": "magenta",
                    "cover_letter_mediator_final": "green",
                    "application_submitted": "bold green",
                    "application_failed": "bold red",
                }
                color = type_colors.get(event_type, "")

                console.print(f"\n[{color}]--- {event_type} ({created}) ---[/{color}]")

                # Show metadata highlights
                meta = event.get("metadata", {})
                if meta.get("feedback"):
                    console.print(f"  [italic]Feedback: {meta['feedback']}[/italic]")
                if meta.get("changes_made"):
                    console.print("  Changes:")
                    for change in meta["changes_made"][:5]:
                        console.print(f"    - {change}")
                if meta.get("overall_score"):
                    console.print(f"  Score: {meta['overall_score']}")

                # Show content preview
                content = event.get("content", "")
                if content:
                    preview = content[:300]
                    if len(content) > 300:
                        preview += "..."
                    console.print(f"  Content: {preview}")

            console.print(f"\n[dim]Full audit files: {config.application.output_dir}/audit/{job_id}/[/dim]")

        asyncio.run(_show())
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def review(config_path: str) -> None:
    """Interactively review pending applications. Approve or skip each one."""
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)

        async def _review():
            apps = await db.get_applications(
                status=ApplicationStatus("pending"), limit=100
            )

            if not apps:
                console.print("[yellow]No pending applications to review.[/yellow]")
                console.print("Run 'super-job-apply run --dry-run' first to discover and tailor applications.")
                return

            console.print(f"\n[bold]Reviewing {len(apps)} pending applications[/bold]\n")

            approved_count = 0
            skipped_count = 0

            for i, app in enumerate(apps, 1):
                console.print("=" * 60)
                console.print(
                    f"[bold]({i}/{len(apps)}) {app.get('company_name', '?')} — "
                    f"{app.get('job_title', '?')}[/bold]"
                )
                console.print(f"  URL: {app.get('careers_url', '—')}")
                console.print(
                    f"  Match score: [{'green' if (app.get('match_score') or 0) >= 0.7 else 'yellow'}]"
                    f"{app.get('match_score', 0):.2f}[/]"
                )

                # Show location/work type
                location = app.get("location") or "Not specified"
                work_type = app.get("work_type") or "Not specified"
                console.print(f"  Location: {location} | Type: {work_type}")

                # Show tailored materials
                if app.get("resume_path"):
                    console.print(f"  Resume: [green]{app['resume_path']}[/green]")
                if app.get("cover_letter_path"):
                    console.print(f"  Cover letter: [green]{app['cover_letter_path']}[/green]")

                # Show cover letter preview from the .docx if it exists
                cl_path = app.get("cover_letter_path")
                if cl_path:
                    try:
                        from pathlib import Path
                        if Path(cl_path).exists():
                            from docx import Document as DocxDocument
                            doc = DocxDocument(cl_path)
                            cl_text = "\n".join(
                                p.text for p in doc.paragraphs if p.text.strip()
                            )
                            # Show first ~500 chars
                            preview = cl_text[:500]
                            if len(cl_text) > 500:
                                preview += "..."
                            console.print(f"\n  [dim]Cover letter preview:[/dim]")
                            console.print(f"  [italic]{preview}[/italic]")
                    except Exception:
                        pass

                # Show job description preview
                desc = app.get("full_description", "")
                if desc:
                    preview = desc[:300]
                    if len(desc) > 300:
                        preview += "..."
                    console.print(f"\n  [dim]Job description preview:[/dim]")
                    console.print(f"  {preview}")

                console.print()

                # Ask for approval
                choice = click.prompt(
                    "  [y]es / [n]o / [a]pprove all remaining / [q]uit",
                    type=click.Choice(["y", "n", "a", "q"], case_sensitive=False),
                    default="y",
                )

                if choice == "q":
                    console.print("[yellow]Review stopped.[/yellow]")
                    break
                elif choice == "a":
                    # Approve this one and all remaining
                    await db.update_application(
                        app["id"], status=ApplicationStatus.APPROVED
                    )
                    approved_count += 1
                    console.print("  [green]APPROVED[/green]")
                    # Approve all remaining
                    remaining = apps[i:]  # apps after current (i is 0-indexed from enumerate starting at 1)
                    for rem_app in remaining:
                        await db.update_application(
                            rem_app["id"], status=ApplicationStatus.APPROVED
                        )
                        approved_count += 1
                    console.print(
                        f"\n  [bold green]Approved all {len(remaining) + 1} remaining applications[/bold green]"
                    )
                    break
                elif choice == "y":
                    await db.update_application(
                        app["id"], status=ApplicationStatus.APPROVED
                    )
                    approved_count += 1
                    console.print("  [green]APPROVED[/green]")
                else:
                    await db.update_application(
                        app["id"], status=ApplicationStatus.SKIPPED
                    )
                    skipped_count += 1
                    console.print("  [dim]Skipped[/dim]")

                console.print()

            console.print("=" * 60)
            console.print(
                f"[bold]Review complete:[/bold] "
                f"[green]{approved_count} approved[/green], "
                f"[dim]{skipped_count} skipped[/dim]"
            )
            if approved_count > 0:
                console.print(
                    "\nRun [bold]super-job-apply submit[/bold] to apply to approved jobs."
                )

        asyncio.run(_review())
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command()
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def submit(config_path: str) -> None:
    """Submit all approved applications (fills forms and clicks submit).

    Uses concurrent browser sessions based on max_concurrent_browsers setting.
    """
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)
        settings = config.application
        candidate = config.candidate

        from .applicator.browser import apply_with_retry
        from .applicator.email_verifier import EmailWatcher
        from .audit import AuditTrail

        audit_trail = AuditTrail(
            db_path=settings.db_path, output_dir=settings.output_dir
        )

        MAX_TOTAL_ATTEMPTS = 3

        async def _submit_one(app: dict, index: int, total: int, email_watcher=None) -> dict:
            """Submit a single application and record the result."""
            import json
            from datetime import datetime, timezone
            from pathlib import Path

            from .models import JobPosting

            company = app.get("company_name", "?")
            title = app.get("job_title", "?")

            # Check retry count — skip if already attempted too many times
            prior_attempts = app.get("retry_count", 0) or 0
            if prior_attempts >= MAX_TOTAL_ATTEMPTS:
                console.print(
                    f"[dim]({index}/{total}) Skipping {company} — {title} "
                    f"(already attempted {prior_attempts} times)[/dim]"
                )
                await db.update_application(
                    app["id"],
                    status=ApplicationStatus.SKIPPED,
                    error_message=f"Exceeded {MAX_TOTAL_ATTEMPTS} attempt limit",
                )
                return {"success": False, "message": "Max attempts exceeded", "company": company, "job_title": title}

            console.print(
                f"[bold]({index}/{total}) Submitting: {company} — {title}[/bold]"
            )

            job = JobPosting(
                id=app["job_id"],
                company_name=company,
                job_title=title,
                careers_url=app.get("careers_url", ""),
                location=app.get("location"),
                work_type=app.get("work_type"),
                requirements=json.loads(app["requirements"]) if app.get("requirements") else [],
                responsibilities=json.loads(app["responsibilities"]) if app.get("responsibilities") else [],
                full_description=app.get("full_description", ""),
            )

            cover_letter_text = None
            cl_path = app.get("cover_letter_path")
            if cl_path:
                try:
                    if Path(cl_path).exists():
                        from docx import Document as DocxDocument
                        doc = DocxDocument(cl_path)
                        cover_letter_text = "\n".join(
                            p.text for p in doc.paragraphs if p.text.strip()
                        )
                except Exception:
                    pass

            result = await apply_with_retry(
                job, candidate, settings,
                resume_path=app.get("resume_path"),
                cover_letter_text=cover_letter_text,
                submit=True,
                email_watcher=email_watcher,
            )

            success = result.get("success", False)
            new_status = ApplicationStatus.APPLIED if success else ApplicationStatus.FAILED
            await db.update_application(
                app["id"],
                status=new_status,
                session_url=result.get("session_url"),
                error_message=result.get("message") if not success else None,
                applied_at=datetime.now(timezone.utc) if success else None,
                retry_count=prior_attempts + 1,
            )

            await audit_trail.log_application_result(
                job_id=app["job_id"],
                application_id=app["id"],
                success=success,
                session_url=result.get("session_url"),
                error=result.get("message") if not success else None,
            )

            # Log account creation if it happened
            if result.get("account_created"):
                from .audit import AuditEvent, AuditEventType
                await audit_trail.record(AuditEvent(
                    job_id=app["job_id"],
                    application_id=app["id"],
                    event_type=AuditEventType.ACCOUNT_CREATED,
                    content=f"Account created on {company} job site",
                    metadata={
                        "email": candidate.email,
                        "site_url": app.get("careers_url", ""),
                        "company": company,
                    },
                ))
                console.print(f"  [cyan]Account created on {company}[/cyan]")

            status_str = "[green]SUBMITTED[/green]" if success else "[red]FAILED[/red]"
            console.print(f"  {status_str} {company} — {title}")
            if result.get("session_url"):
                console.print(f"    Session: {result['session_url']}")
            if not success and result.get("message"):
                console.print(f"    Error: {result['message'][:150]}")

            result["company"] = company
            result["job_title"] = title
            return result

        async def _submit():
            await audit_trail.init()

            # Start email watcher for security code handling
            email_watcher = EmailWatcher()
            email_watcher.start()
            console.print(
                f"[dim]Email watcher: {'active' if email_watcher.available else 'disabled'}"
                f" ({email_watcher.imap_email})[/dim]"
            )

            apps = await db.get_applications(
                status=ApplicationStatus("approved"), limit=500
            )

            if not apps:
                console.print("[yellow]No approved applications to submit.[/yellow]")
                console.print("Run 'super-job-apply review' first to approve applications.")
                return

            # Filter out blocked ATS platforms before launching browsers
            from .ats_detection import BLOCKED_ATS, detect_ats_platform
            original_count = len(apps)
            blocked_apps = []
            submittable_apps = []
            for app in apps:
                ats = detect_ats_platform(app.get("careers_url", ""))
                if ats in BLOCKED_ATS:
                    blocked_apps.append((app, ats))
                else:
                    submittable_apps.append(app)

            if blocked_apps:
                console.print(
                    f"[yellow]Skipping {len(blocked_apps)} jobs on blocked ATS "
                    f"({', '.join(sorted(BLOCKED_ATS))}):[/yellow]"
                )
                for app, ats in blocked_apps:
                    console.print(
                        f"  [dim]{app.get('company_name', '?')} — {app.get('job_title', '?')} ({ats})[/dim]"
                    )
                    await db.update_application(
                        app["id"],
                        status=ApplicationStatus.SKIPPED,
                        error_message=f"Blocked ATS: {ats}",
                    )

            apps = submittable_apps
            if not apps:
                console.print("[yellow]All approved applications are on blocked ATS platforms.[/yellow]")
                console.print("Run 'super-job-apply submittable' to see available platforms.")
                return

            max_browsers = settings.max_concurrent_browsers
            mode = f"concurrent, {max_browsers} browsers" if settings.concurrent else "sequential"
            console.print(
                f"\n[bold]Submitting {len(apps)} approved applications ({mode})...[/bold]\n"
            )

            results = []
            if settings.concurrent and max_browsers > 1:
                # Process in chunks with delay between batches to avoid API rate limits
                for i in range(0, len(apps), max_browsers):
                    chunk = apps[i : i + max_browsers]
                    batch_num = i // max_browsers + 1
                    console.print(
                        f"\n[dim]--- Batch {batch_num} "
                        f"({len(chunk)} parallel sessions) ---[/dim]"
                    )
                    chunk_results = await asyncio.gather(
                        *[
                            _submit_one(app, i + j + 1, len(apps), email_watcher=email_watcher)
                            for j, app in enumerate(chunk)
                        ]
                    )
                    results.extend(chunk_results)
                    # Brief pause between batches to avoid API rate limits
                    if i + max_browsers < len(apps):
                        console.print("[dim]  Pausing 30s between batches (rate limit)...[/dim]")
                        await asyncio.sleep(30)
            else:
                for i, app in enumerate(apps):
                    result = await _submit_one(app, i + 1, len(apps), email_watcher=email_watcher)
                    results.append(result)

            # Stop email watcher
            email_watcher.stop()

            # Summary
            ok = sum(1 for r in results if r.get("success"))
            fail = len(results) - ok
            confirmed = sum(1 for r in results if r.get("confirmed"))
            fields = sum(r.get("fields_filled", 0) for r in results)
            console.print("\n" + "=" * 60)
            console.print(
                f"[bold]Submission complete:[/bold] "
                f"[green]{ok} submitted[/green], "
                f"[bold green]{confirmed} confirmed[/bold green], "
                f"[red]{fail} failed[/red], "
                f"{fields} total fields filled"
            )

        asyncio.run(_submit())
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command(name="retry-failed")
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def retry_failed(config_path: str) -> None:
    """Re-submit all failed applications. Moves them back to approved and runs submit."""
    try:
        config = load_config(config_path)
        db = Database(config.application.db_path)
        settings = config.application
        candidate = config.candidate

        from .applicator.browser import apply_with_retry
        from .audit import AuditTrail

        audit_trail = AuditTrail(
            db_path=settings.db_path, output_dir=settings.output_dir
        )

        async def _retry():
            await audit_trail.init()

            # Get all failed applications
            failed_apps = await db.get_applications(
                status=ApplicationStatus("failed"), limit=500
            )

            if not failed_apps:
                console.print("[yellow]No failed applications to retry.[/yellow]")
                return

            console.print(
                f"\n[bold]Retrying {len(failed_apps)} failed applications...[/bold]"
            )

            # Show what failed and why
            for i, app in enumerate(failed_apps, 1):
                err = app.get("error_message", "Unknown error")
                console.print(
                    f"  {i}. {app.get('company_name', '?')} — {app.get('job_title', '?')}"
                )
                console.print(f"     [dim]Previous error: {err[:120]}[/dim]")

            console.print()
            if not click.confirm(f"Retry all {len(failed_apps)} failed applications?"):
                return

            # Reset to approved status and increment retry count
            for app in failed_apps:
                await db.update_application(
                    app["id"],
                    status=ApplicationStatus.APPROVED,
                    error_message=None,
                )

            # Now run the same submit logic
            apps = await db.get_applications(
                status=ApplicationStatus("approved"), limit=500
            )

            max_browsers = settings.max_concurrent_browsers
            mode = f"concurrent, {max_browsers} browsers" if settings.concurrent else "sequential"
            console.print(
                f"\n[bold]Submitting {len(apps)} applications ({mode})...[/bold]\n"
            )

            # Reuse _submit_one from the submit command
            import json
            from datetime import datetime, timezone
            from pathlib import Path

            from .audit import AuditEvent, AuditEventType
            from .models import JobPosting

            async def _submit_one(app_data, idx, total):
                company = app_data.get("company_name", "?")
                title = app_data.get("job_title", "?")

                prior_attempts = app_data.get("retry_count", 0) or 0
                if prior_attempts >= 3:
                    console.print(
                        f"[dim]({idx}/{total}) Skipping {company} — {title} "
                        f"(already attempted {prior_attempts} times)[/dim]"
                    )
                    await db.update_application(
                        app_data["id"],
                        status=ApplicationStatus.SKIPPED,
                        error_message="Exceeded 3 attempt limit",
                    )
                    return {"success": False, "message": "Max attempts exceeded", "company": company, "job_title": title}

                console.print(
                    f"[bold]({idx}/{total}) Retrying: {company} — {title}[/bold]"
                )

                job = JobPosting(
                    id=app_data["job_id"],
                    company_name=company,
                    job_title=title,
                    careers_url=app_data.get("careers_url", ""),
                    location=app_data.get("location"),
                    work_type=app_data.get("work_type"),
                    requirements=json.loads(app_data["requirements"]) if app_data.get("requirements") else [],
                    responsibilities=json.loads(app_data["responsibilities"]) if app_data.get("responsibilities") else [],
                    full_description=app_data.get("full_description", ""),
                )

                cover_letter_text = None
                cl_path = app_data.get("cover_letter_path")
                if cl_path:
                    try:
                        if Path(cl_path).exists():
                            from docx import Document as DocxDocument
                            doc = DocxDocument(cl_path)
                            cover_letter_text = "\n".join(
                                p.text for p in doc.paragraphs if p.text.strip()
                            )
                    except Exception:
                        pass

                result = await apply_with_retry(
                    job, candidate, settings,
                    resume_path=app_data.get("resume_path"),
                    cover_letter_text=cover_letter_text,
                    submit=True,
                )

                success = result.get("success", False)
                new_status = ApplicationStatus.APPLIED if success else ApplicationStatus.FAILED
                await db.update_application(
                    app_data["id"],
                    status=new_status,
                    session_url=result.get("session_url"),
                    error_message=result.get("message") if not success else None,
                    applied_at=datetime.now(timezone.utc) if success else None,
                    retry_count=prior_attempts + 1,
                )

                await audit_trail.log_application_result(
                    job_id=app_data["job_id"],
                    application_id=app_data["id"],
                    success=success,
                    session_url=result.get("session_url"),
                    error=result.get("message") if not success else None,
                )

                if result.get("account_created"):
                    await audit_trail.record(AuditEvent(
                        job_id=app_data["job_id"],
                        application_id=app_data["id"],
                        event_type=AuditEventType.ACCOUNT_CREATED,
                        content=f"Account created on {company} job site",
                        metadata={
                            "email": candidate.email,
                            "site_url": app_data.get("careers_url", ""),
                            "company": company,
                        },
                    ))
                    console.print(f"  [cyan]Account created on {company}[/cyan]")

                status_str = "[green]SUBMITTED[/green]" if success else "[red]FAILED[/red]"
                console.print(f"  {status_str} {company} — {title}")
                if result.get("session_url"):
                    console.print(f"    Session: {result['session_url']}")
                if not success and result.get("message"):
                    console.print(f"    Error: {result['message'][:150]}")

                result["company"] = company
                result["job_title"] = title
                return result

            results = []
            if settings.concurrent and max_browsers > 1:
                for i in range(0, len(apps), max_browsers):
                    chunk = apps[i : i + max_browsers]
                    console.print(
                        f"\n[dim]--- Batch {i // max_browsers + 1} "
                        f"({len(chunk)} parallel sessions) ---[/dim]"
                    )
                    chunk_results = await asyncio.gather(
                        *[
                            _submit_one(app, i + j + 1, len(apps))
                            for j, app in enumerate(chunk)
                        ]
                    )
                    results.extend(chunk_results)
            else:
                for i, app in enumerate(apps):
                    result = await _submit_one(app, i + 1, len(apps))
                    results.append(result)

            ok = sum(1 for r in results if r.get("success"))
            fail = len(results) - ok
            console.print("\n" + "=" * 60)
            console.print(
                f"[bold]Retry complete:[/bold] "
                f"[green]{ok} submitted[/green], "
                f"[red]{fail} still failed[/red]"
            )

        asyncio.run(_retry())
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


@main.command(name="resolve-urls")
@click.option("--config", "config_path", default="config.yaml", help="Path to config YAML")
def resolve_urls(config_path: str) -> None:
    """Resolve aggregator URLs to direct company career pages.

    Finds the actual company application page for jobs that were discovered
    via Indeed, Dice, and other aggregator sites. Resolved jobs are reset
    to 'approved' for resubmission.
    """
    try:
        config = load_config(config_path)

        from .discovery.url_resolver import resolve_aggregator_jobs

        async def _resolve():
            result = await resolve_aggregator_jobs(config.application.db_path)
            console.print(f"\n[bold]URL Resolution Complete[/bold]")
            console.print(f"  [green]Resolved: {result['resolved']}[/green] (reset to approved)")
            console.print(f"  [yellow]Unresolved: {result['unresolved']}[/yellow]")
            console.print(f"  [dim]Skipped (generic): {result['skipped']}[/dim]")

            if result["resolved"] > 0:
                console.print(
                    f"\nRun [bold]super-job-apply submit[/bold] to apply to the {result['resolved']} resolved jobs."
                )

        asyncio.run(_resolve())
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


if __name__ == "__main__":
    main()
